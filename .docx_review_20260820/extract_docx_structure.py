from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


source = Path(r"C:\Users\53507\Desktop\计算机本科 AI科研与香港申请长期课表_2026-2029.docx")
document = Document(source)

print(f"TITLE: {document.core_properties.title}")
print(f"SUBJECT: {document.core_properties.subject}")
print(f"AUTHOR: {document.core_properties.author}")
print("=== BODY IN DOCUMENT ORDER ===")

paragraph_number = 0
table_number = 0

for child in document.element.body.iterchildren():
    if child.tag == qn("w:p"):
        paragraph_number += 1
        paragraph = Paragraph(child, document)
        text = paragraph.text.strip()
        if text:
            print(f"[P{paragraph_number} | {paragraph.style.name}] {text}")
    elif child.tag == qn("w:tbl"):
        table_number += 1
        table = Table(child, document)
        print(f"[TABLE {table_number} | {len(table.rows)} rows x {len(table.columns)} cols]")
        for row_number, row in enumerate(table.rows, start=1):
            values = [" / ".join(part.strip() for part in cell.text.splitlines() if part.strip()) for cell in row.cells]
            print(f"  R{row_number}: " + " || ".join(values))

print("=== HEADERS AND FOOTERS ===")
for section_number, section in enumerate(document.sections, start=1):
    header_text = " | ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
    footer_text = " | ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
    print(f"SECTION {section_number} HEADER: {header_text}")
    print(f"SECTION {section_number} FOOTER: {footer_text}")

print(f"=== TOTALS: {paragraph_number} paragraphs, {table_number} tables, {len(document.sections)} sections ===")
